import os
import docx
import docx2txt
from docx.shared import RGBColor

def create_products_folder():
    """
    Create a 'products' folder if it doesn't exist
    """
    if not os.path.exists("products"):
        os.makedirs("products")

def read_docx(file_path):
    """
    Read the contents of a .docx file
    """
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_doc(file_path):
    """
    Read the contents of a .doc file
    """
    text = docx2txt.process(file_path)
    return text

def detect_highlight_colors(text):
    """
    Detect the highlight colors used in the text
    """
    # Your logic to detect highlight colors goes here
    # This can involve analyzing the formatting of the text
    
    # For demonstration purposes, let's assume we have detected some colors
    colors = [RGBColor(255, 0, 0), RGBColor(0, 255, 0)]  # Example colors
    return colors

def apply_highlighting(doc, colors):
    """
    Apply highlighting to the document based on the detected colors
    """
    for color in colors:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb == color:
                    run.font.highlight_color = color

def main(file_path):
    """
    Main function to process the document
    """
    if file_path.endswith('.docx'):
        text = read_docx(file_path)
    elif file_path.endswith('.doc'):
        text = read_doc(file_path)
    else:
        print("Unsupported file format")
        return

    colors = detect_highlight_colors(text)

    doc = docx.Document()
    doc.add_paragraph(text)
    apply_highlighting(doc, colors)

    create_products_folder()
    doc_path = "products/highlighted_document.docx"
    doc.save(doc_path)
    print(f"Resultant document saved to {doc_path}")

if __name__ == "__main__":
    file_path = input("Enter the path to the .doc or .docx file: ")
    main(file_path)